"""
inject_into_template.py — 将转换后内容注入用户提供的 Word 模板
用法: python inject_into_template.py template.docx content.docx output.docx [标记文本]

模板中应包含一个标记段落（默认 "正文开始"），脚本会在该位置插入内容并删除标记。
如果找不到标记，内容追加到模板末尾。
"""
import sys
import copy
from docx import Document


def inject(template_path, content_path, output_path,
           insert_marker='正文开始'):
    """将 content_path 的全部内容注入 template_path 的指定位置"""
    template = Document(template_path)
    content = Document(content_path)

    body = template.element.body
    body_children = list(body)

    # 找到标记段落
    insert_idx = len(body_children)  # 默认追加到末尾
    marker_found = False
    for i, para in enumerate(template.paragraphs):
        if insert_marker in para.text:
            insert_idx = body_children.index(para._p)
            para._p.getparent().remove(para._p)
            marker_found = True
            print(f'找到标记 "{insert_marker}"，在位置 {insert_idx} 插入内容')
            break

    if not marker_found:
        print(f'未找到标记 "{insert_marker}"，内容将追加到文档末尾')

    # 复制内容文档的所有元素
    content_elements = list(content.element.body)
    for j, elem in enumerate(content_elements):
        new_elem = copy.deepcopy(elem)
        body.insert(insert_idx + j, new_elem)

    template.save(output_path)
    print(f'模板填充完成 → {output_path}')
    print(f'共注入 {len(content_elements)} 个元素')


if __name__ == '__main__':
    if len(sys.argv) < 4:
        print('用法: python inject_into_template.py '
              '<template.docx> <content.docx> <output.docx> [标记文本]')
        sys.exit(1)

    tpl = sys.argv[1]
    cnt = sys.argv[2]
    out = sys.argv[3]
    marker = sys.argv[4] if len(sys.argv) > 4 else '正文开始'
    inject(tpl, cnt, out, marker)
