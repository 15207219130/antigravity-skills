"""
postprocess_docx.py — Word 文档后处理 v3
功能：
  1. 中文 → 宋体，英文 → Times New Roman，标题 → 黑体
  2. 移除所有斜体
  3. 三线表
  4. 公式编号：识别 EQNUM(N) 标记 → 居中公式 + 右对齐 (N)

用法: python postprocess_docx.py input.docx [output.docx]
"""
import re
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree


# ============================================================
# 字体
# ============================================================
def set_run_fonts(run, cn='宋体', en='Times New Roman', sz=12):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:cs'), en)
    run.font.size = Pt(sz)


def fix_fonts(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            set_run_fonts(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_fonts(run)


def fix_heading_fonts(doc):
    sz_map = {'Heading 1': 16, 'Heading 2': 14, 'Heading 3': 13}
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            sz = sz_map.get(para.style.name, 13)
            for run in para.runs:
                set_run_fonts(run, cn='黑体', sz=sz)


# ============================================================
# 移除斜体
# ============================================================
def remove_all_italics(doc):
    count = 0
    all_runs = []
    for para in doc.paragraphs:
        all_runs.extend(para.runs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_runs.extend(para.runs)
    for run in all_runs:
        rPr = run._r.find(qn('w:rPr'))
        if rPr is not None:
            for tag in ['w:i', 'w:iCs']:
                for elem in rPr.findall(qn(tag)):
                    rPr.remove(elem)
                    count += 1
        if run.italic:
            run.italic = False
            count += 1
    print(f'  移除了 {count} 处斜体属性')


# ============================================================
# 三线表
# ============================================================
def apply_three_line_table(table):
    THICK = {"sz": "12", "val": "single", "color": "000000"}
    NONE = {"val": "nil"}
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
            # 清除旧边框
            for old in list(tcBorders):
                tcBorders.remove(old)
            if i == 0:
                for edge, style in [('top', THICK), ('bottom', THICK),
                                     ('left', NONE), ('right', NONE)]:
                    e = etree.SubElement(tcBorders, qn(f'w:{edge}'))
                    for k, v in style.items():
                        e.set(qn(f'w:{k}'), str(v))
            elif i == len(table.rows) - 1:
                for edge, style in [('top', NONE), ('bottom', THICK),
                                     ('left', NONE), ('right', NONE)]:
                    e = etree.SubElement(tcBorders, qn(f'w:{edge}'))
                    for k, v in style.items():
                        e.set(qn(f'w:{k}'), str(v))
            else:
                for edge in ['top', 'bottom', 'left', 'right']:
                    e = etree.SubElement(tcBorders, qn(f'w:{edge}'))
                    e.set(qn('w:val'), 'nil')
    # 清表级边框
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tb = tblPr.find(qn('w:tblBorders'))
        if tb is not None:
            tblPr.remove(tb)


def fix_tables(doc):
    for table in doc.tables:
        apply_three_line_table(table)
    print(f'  {len(doc.tables)} 个表格 → 三线表')


# ============================================================
# 公式编号：EQNUM(N) → 居中公式 + 右对齐 (N)
# ============================================================
def fix_equation_numbers(doc):
    """
    找到 EQNUM(N) 段落，将编号合并到上方公式段落。
    使用 Word 经典三栏制表位布局：
      [居中tab] 公式 [右对齐tab] (N)
    
    关键：在 OMML 公式元素前插入一个 tab run 将公式推到居中制表位。
    """
    NSMAP_M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    paragraphs = list(doc.paragraphs)
    to_remove = []
    count = 0

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        m = re.match(r'^EQNUM\((\d+)\)$', text)
        if not m:
            continue

        eq_num = m.group(1)

        if i > 0:
            eq_para = paragraphs[i - 1]
            eq_p = eq_para._p

            # --- 1) 在 OMML 公式前插入一个 tab run（推到居中位置）---
            omath = eq_p.find(f'{{{NSMAP_M}}}oMathPara')
            if omath is None:
                omath = eq_p.find(f'{{{NSMAP_M}}}oMath')

            if omath is not None:
                # 创建一个包含 tab 字符的 w:r 元素
                tab_r = etree.SubElement(eq_p, qn('w:r'))
                tab_elem = etree.SubElement(tab_r, qn('w:tab'))
                # 插入到 OMML 元素之前
                eq_p.insert(list(eq_p).index(omath), tab_r)

            # --- 2) 在末尾添加 tab + (N) ---
            tab_run2 = eq_para.add_run('\t')
            num_run = eq_para.add_run(f'({eq_num})')
            set_run_fonts(num_run)

            # --- 3) 设置制表位：居中 + 右对齐 ---
            pPr = eq_p.get_or_add_pPr()
            tabs = pPr.find(qn('w:tabs'))
            if tabs is None:
                tabs = etree.SubElement(pPr, qn('w:tabs'))
            for old in list(tabs):
                tabs.remove(old)

            # 居中制表位（公式居中）
            tc = etree.SubElement(tabs, qn('w:tab'))
            tc.set(qn('w:val'), 'center')
            tc.set(qn('w:pos'), '4320')  # ~7.6cm, 页面中间

            # 右对齐制表位（编号靠右）
            tr = etree.SubElement(tabs, qn('w:tab'))
            tr.set(qn('w:val'), 'right')
            tr.set(qn('w:pos'), '8640')  # ~15.2cm, 页面右边

            count += 1

        to_remove.append(para._p)

    for p_elem in to_remove:
        p_elem.getparent().remove(p_elem)

    print(f'  {count} 个公式：居中 + 编号右对齐')
    print(f'  删除了 {len(to_remove)} 个 EQNUM 标记段落')


# ============================================================
# Main
# ============================================================
def postprocess(input_docx, output_docx=None):
    if output_docx is None:
        output_docx = input_docx.replace('.docx', '_final.docx')

    doc = Document(input_docx)

    print('[1/5] 字体 (宋体 + TNR)...')
    fix_fonts(doc)

    print('[2/5] 标题字体 (黑体)...')
    fix_heading_fonts(doc)

    print('[3/5] 去斜体...')
    remove_all_italics(doc)

    print('[4/5] 三线表...')
    fix_tables(doc)

    print('[5/5] 公式编号 (居中+右对齐)...')
    fix_equation_numbers(doc)

    doc.save(output_docx)
    print(f'\n完成 → {output_docx}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python postprocess_docx.py <input.docx> [output.docx]')
        sys.exit(1)
    postprocess(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
